import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
import cv2

def scaling_closest_neighbour(img, scale_multiplier):
    height, width = img.shape[:2]

    new_height = int(np.ceil(height * scale_multiplier))
    new_width = int(np.ceil(width * scale_multiplier))

    new_x = np.linspace(0, width - 1, new_width)
    new_y = np.linspace(0, height - 1, new_height)

    if len(img.shape) < 3:
        scaled_img = np.zeros((new_height, new_width), dtype=np.uint8)
    else:
        scaled_img = np.zeros((new_height, new_width, img.shape[2]), dtype=np.uint8)

    for y in range(new_height):
        for x in range(new_width):
            nearest_y = int(np.round(new_y[y]))
            nearest_x = int(np.round(new_x[x]))

            scaled_img[y, x] = img[nearest_y, nearest_x]

    return scaled_img

def scaling_bilinear_interpolation(img, scale_multiplier):
    height, width = img.shape[:2]

    new_height = int(np.ceil(height * scale_multiplier))
    new_width = int(np.ceil(width * scale_multiplier))

    new_x = np.linspace(0, width - 1, new_width)
    new_y = np.linspace(0, height - 1, new_height)

    if len(img.shape) < 3:
        scaled_img = np.zeros((new_height, new_width), dtype=np.float32)
    else:
        scaled_img = np.zeros((new_height, new_width, img.shape[2]), dtype=np.float32)

    for y in range(new_height):
        for x in range(new_width):
            x_orig = new_x[x]
            y_orig = new_y[y]

            x1 = int(np.floor(x_orig))
            x2 = min(x1 + 1, width - 1)
            y1 = int(np.floor(y_orig))
            y2 = min(y1 + 1, height - 1)

            dx = x_orig - x1
            dy = y_orig - y1

            Q11 = img[y1, x1]
            Q21 = img[y1, x2]
            Q12 = img[y2, x1]
            Q22 = img[y2, x2]

            interpolated_value = (
                Q11 * (1 - dx) * (1 - dy) +
                Q21 * dx * (1 - dy) +
                Q12 * (1 - dx) * dy +
                Q22 * dx * dy
            )

            scaled_img[y, x] = interpolated_value

    return np.clip(scaled_img, 0, 255).astype(np.uint8)

# test skalowania w gore

# img= np.zeros((20,20,3),dtype=np.uint8)
# img[1,1,:]=255
#
# scaled_img_closest = scaling_closest_neighbour(img, 10)
# scaled_img_bilinear = scaling_bilinear_interpolation(img, 10)
#
# plt.imshow(img)
# plt.show()
#
# plt.imshow(scaled_img_closest)
# plt.show()
#
# plt.imshow(scaled_img_bilinear)
# plt.show()

# test skalowania w dol

# img = cv2.imread('IMG_BIG/BIG_0001.jpg')
# img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
#
# scaled_img_closest = scaling_closest_neighbour(img, 0.01)
# scaled_img_bilinear = scaling_bilinear_interpolation(img, 0.01)
#
# plt.imshow(img)
# plt.show()
#
# plt.imshow(scaled_img_closest)
# plt.show()
#
# plt.imshow(scaled_img_bilinear)
# plt.show()

def downscale_mean(img, scale_multiplier):
    height, width = img.shape[:2]

    new_height = int(np.ceil(height * scale_multiplier))
    new_width = int(np.ceil(width * scale_multiplier))

    if len(img.shape) < 3:
        downscaled_img = np.zeros((new_height, new_width), dtype=np.float32)
    else:
        downscaled_img = np.zeros((new_height, new_width, img.shape[2]), dtype=np.float32)

    xx = np.linspace(0, height, new_height + 1)
    yy = np.linspace(0, width, new_width + 1)

    for y in range(new_height):
        for x in range(new_width):
            if y > 0:
                y1 = -(xx[y] - xx[y - 1]) / 2
            else:
                y1 = 0
            if y < new_height - 1:
                y2 = (xx[y + 1] - xx[y]) / 2 + 1
            else:
                y2 = 0

            if x > 0:
                x1 = -(yy[x] - yy[x - 1]) / 2
            else:
                x1 = 0
            if x < new_width - 1:
                x2 = (yy[x + 1] - yy[x]) / 2 + 1
            else:
                x2 = 0

            iy = np.round(xx[y] + np.arange(y1, y2)).astype(int)
            ix = np.round(yy[x] + np.arange(x1, x2)).astype(int)

            iy = iy.clip(0, height - 1)
            ix = ix.clip(0, width - 1)

            region = img[iy[0]:iy[-1] + 1, ix[0]:ix[-1] + 1]

            if len(region.shape) == 3:
                downscaled_img[y, x] = np.mean(region, axis=(0, 1))
            else:
                downscaled_img[y, x] = np.mean(region)

    downscaled_img = np.clip(downscaled_img, 0, 255)

    return downscaled_img.astype(np.uint8)

# img = cv2.imread('IMG_BIG/BIG_0001.jpg')
# img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
#
# downscaled_img_mean = downscale_mean(img, 0.01)
#
# plt.imshow(img)
# plt.show()
#
# plt.imshow(downscaled_img_mean)
# plt.show()

def downscale_weighted_mean(img, scale_multiplier):
    height, width = img.shape[:2]

    new_height = int(np.ceil(height * scale_multiplier))
    new_width = int(np.ceil(width * scale_multiplier))

    if len(img.shape) < 3:
        downscaled_img = np.zeros((new_height, new_width), dtype=np.float32)
    else:
        downscaled_img = np.zeros((new_height, new_width, img.shape[2]), dtype=np.float32)

    xx = np.linspace(0, height, new_height + 1)
    yy = np.linspace(0, width, new_width + 1)

    for y in range(new_height):
        for x in range(new_width):
            if y > 0:
                y1 = -(xx[y] - xx[y - 1]) / 2
            else:
                y1 = 0
            if y < new_height - 1:
                y2 = (xx[y + 1] - xx[y]) / 2 + 1
            else:
                y2 = 0

            if x > 0:
                x1 = -(yy[x] - yy[x - 1]) / 2
            else:
                x1 = 0
            if x < new_width - 1:
                x2 = (yy[x + 1] - yy[x]) / 2 + 1
            else:
                x2 = 0

            iy = np.round(xx[y] + np.arange(y1, y2)).astype(int)
            ix = np.round(yy[x] + np.arange(x1, x2)).astype(int)

            iy = iy.clip(0, height - 1)
            ix = ix.clip(0, width - 1)

            region = img[iy[0]:iy[-1] + 1, ix[0]:ix[-1] + 1]

            weights = np.random.rand(region.shape[0], region.shape[1])

            weights_sum = np.sum(weights)
            weights /= weights_sum

            if len(region.shape) == 3:
                weighted_avg = np.sum(region * weights[:, :, np.newaxis], axis=(0, 1))
            else:
                weighted_avg = np.sum(region * weights, axis=(0, 1))

            downscaled_img[y, x] = weighted_avg

    downscaled_img = np.clip(downscaled_img, 0, 255)

    return downscaled_img.astype(np.uint8)

# img = cv2.imread('IMG_BIG/BIG_0001.jpg')
# img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
#
# downscaled_img_weighted_mean = downscale_weighted_mean(img, 0.01)
#
# plt.imshow(img)
# plt.show()
#
# plt.imshow(downscaled_img_weighted_mean)
# plt.show()

def downscale_median(img, scale_multiplier):
    height, width = img.shape[:2]

    new_height = int(np.ceil(height * scale_multiplier))
    new_width = int(np.ceil(width * scale_multiplier))

    if len(img.shape) < 3:
        downscaled_img = np.zeros((new_height, new_width), dtype=np.float32)
    else:
        downscaled_img = np.zeros((new_height, new_width, img.shape[2]), dtype=np.float32)

    xx = np.linspace(0, height, new_height + 1)
    yy = np.linspace(0, width, new_width + 1)

    for y in range(new_height):
        for x in range(new_width):
            if y > 0:
                y1 = -(xx[y] - xx[y - 1]) / 2
            else:
                y1 = 0
            if y < new_height - 1:
                y2 = (xx[y + 1] - xx[y]) / 2 + 1
            else:
                y2 = 0

            if x > 0:
                x1 = -(yy[x] - yy[x - 1]) / 2
            else:
                x1 = 0
            if x < new_width - 1:
                x2 = (yy[x + 1] - yy[x]) / 2 + 1
            else:
                x2 = 0

            iy = np.round(xx[y] + np.arange(y1, y2)).astype(int)
            ix = np.round(yy[x] + np.arange(x1, x2)).astype(int)

            iy = iy.clip(0, height - 1)
            ix = ix.clip(0, width - 1)

            region = img[iy[0]:iy[-1] + 1, ix[0]:ix[-1] + 1]

            if len(region.shape) == 3:
                downscaled_img[y, x] = np.median(region, axis=(0, 1))
            else:
                downscaled_img[y, x] = np.median(region)

    downscaled_img = np.clip(downscaled_img, 0, 255)

    return downscaled_img.astype(np.uint8)

# img = cv2.imread('IMG_BIG/BIG_0001.jpg')
# img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
#
# downscaled_img_median = downscale_median(img, 0.01)
#
# plt.imshow(img)
# plt.show()
#
# plt.imshow(downscaled_img_median)
# plt.show()

def image_plot_maker(fragment_list):
    fig, ax = plt.subplots(1, len(fragment_list), figsize=(len(fragment_list) * 6, 6))

    if len(fragment_list) == 1:
        ax = [ax]

    for i, fragment in enumerate(fragment_list):
        ax[i].imshow(fragment)
        ax[i].axis('off')
        ax[i].set_xlim(0, fragment.shape[1])
        ax[i].set_ylim(fragment.shape[0], 0)

    memfile = BytesIO()
    plt.savefig(memfile, format='png', bbox_inches='tight', pad_inches=0)
    plt.close(fig)
    memfile.seek(0)
    return memfile

df = pd.DataFrame()

df = pd.DataFrame(
    data={
        'SmallImagesFilename':['IMG_SMALL/SMALL_0004.jpg', 'IMG_SMALL/SMALL_0005.jpg', 'IMG_SMALL/SMALL_0006.jpg'],
        'SmallImageFragments':[[[25,45,40,70],[34,47,65,82]], [[9,68,6,53],[60,90,51,78]], [[36,61,8,30],[31,56,77,99]]],
        'BigImagesFilename':['IMG_BIG/BIG_0001.jpg', 'IMG_BIG/BIG_0002.jpg', 'IMG_BIG/BIG_0003.jpg'],
        'BigImageFragments':[[[2000,2700,1300,2000],[1200,1700,4000,4600]], [[1200,1800,2400,3000],[1700,2200,3900,4500]], [[3300,3800,2200,2700],[600,1100,1600,2100]]]
    })

document = Document()

document.add_heading('Skalowanie x3 małego obrazka', 0)

for index, row in df.iterrows():
    small_img = plt.imread(row['SmallImagesFilename'])

    for f in row['SmallImageFragments']:
        original_fragment = small_img[f[0]:f[1], f[2]:f[3]]

        nearest_neighbor_fragment = scaling_closest_neighbour(original_fragment, 3)
        bilinear_fragment = scaling_bilinear_interpolation(original_fragment, 3)

        memfile = image_plot_maker([original_fragment, nearest_neighbor_fragment, bilinear_fragment])
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('Oryginalny fragment | Metoda najbliższego sąsiada | Metoda interpolacji dwuliniowej')

        original_edges = cv2.Canny(original_fragment, threshold1=100, threshold2=200)
        nearest_neighbor_edges = cv2.Canny(nearest_neighbor_fragment, threshold1=100, threshold2=200)
        bilinear_edges = cv2.Canny(bilinear_fragment, threshold1=100, threshold2=200)

        memfile_edges = image_plot_maker([original_edges, nearest_neighbor_edges, bilinear_edges])
        document.add_picture(memfile_edges, width=Inches(6))
        memfile_edges.close()
        document.add_paragraph('Krawędzie: Oryginalny | Najbliższy sąsiad | Interpolacja dwuliniowa')
        document.add_paragraph('')

document.add_heading('Skalowanie x5 małego obrazka', 0)

for index, row in df.iterrows():
    small_img = plt.imread(row['SmallImagesFilename'])

    for f in row['SmallImageFragments']:
        original_fragment = small_img[f[0]:f[1], f[2]:f[3]]

        nearest_neighbor_fragment = scaling_closest_neighbour(original_fragment, 5)
        bilinear_fragment = scaling_bilinear_interpolation(original_fragment, 5)

        memfile = image_plot_maker([original_fragment, nearest_neighbor_fragment, bilinear_fragment])
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('Oryginalny fragment | Metoda najbliższego sąsiada | Metoda interpolacji dwuliniowej')

        original_edges = cv2.Canny(original_fragment, threshold1=100, threshold2=200)
        nearest_neighbor_edges = cv2.Canny(nearest_neighbor_fragment, threshold1=100, threshold2=200)
        bilinear_edges = cv2.Canny(bilinear_fragment, threshold1=100, threshold2=200)

        memfile_edges = image_plot_maker([original_edges, nearest_neighbor_edges, bilinear_edges])
        document.add_picture(memfile_edges, width=Inches(6))
        memfile_edges.close()
        document.add_paragraph('Krawędzie: Oryginalny | Najbliższy sąsiad | Interpolacja dwuliniowa')
        document.add_paragraph('')

document.add_heading('Skalowanie x0.1 dużego obrazu', 0)

for index, row in df.iterrows():
    big_img = plt.imread(row['BigImagesFilename'])

    for f in row['BigImageFragments']:
        original_fragment = big_img[f[0]:f[1], f[2]:f[3]]

        nearest_neighbor_fragment = scaling_closest_neighbour(original_fragment, 0.1)
        bilinear_fragment = scaling_bilinear_interpolation(original_fragment, 0.1)
        mean_fragment = downscale_mean(original_fragment, 0.1)
        weighted_mean_fragment = downscale_weighted_mean(original_fragment, 0.1)
        median_fragment = downscale_median(original_fragment, 0.1)

        memfile = image_plot_maker([original_fragment, nearest_neighbor_fragment, bilinear_fragment])
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('Oryginalny fragment | Metoda najbliższego sąsiada | Metoda interpolacji dwuliniowej')

        original_edges = cv2.Canny(original_fragment, threshold1=100, threshold2=200)
        nearest_neighbor_edges = cv2.Canny(nearest_neighbor_fragment, threshold1=100, threshold2=200)
        bilinear_edges = cv2.Canny(bilinear_fragment, threshold1=100, threshold2=200)

        memfile_edges = image_plot_maker([original_edges, nearest_neighbor_edges, bilinear_edges])
        document.add_picture(memfile_edges, width=Inches(6))
        memfile_edges.close()
        document.add_paragraph('Krawędzie: Oryginalny | Najbliższy sąsiad | Interpolacja dwuliniowa')

        memfile = image_plot_maker([mean_fragment, weighted_mean_fragment, median_fragment])
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('Średnia | Średnia ważona | Mediana')

        mean_edges = cv2.Canny(mean_fragment, threshold1=100, threshold2=200)
        weighted_mean_edges = cv2.Canny(weighted_mean_fragment, threshold1=100, threshold2=200)
        median_edges = cv2.Canny(median_fragment, threshold1=100, threshold2=200)

        memfile_edges = image_plot_maker([mean_edges, weighted_mean_edges, median_edges])
        document.add_picture(memfile_edges, width=Inches(6))
        memfile_edges.close()
        document.add_paragraph('Krawędzie: Średnia | Średnia ważona | Mediana')
        document.add_paragraph('')

document.add_heading('Skalowanie x0.05 dużego obrazu', 0)

for index, row in df.iterrows():
    big_img = plt.imread(row['BigImagesFilename'])

    for f in row['BigImageFragments']:
        original_fragment = big_img[f[0]:f[1], f[2]:f[3]]

        nearest_neighbor_fragment = scaling_closest_neighbour(original_fragment, 0.05)
        bilinear_fragment = scaling_bilinear_interpolation(original_fragment, 0.05)
        mean_fragment = downscale_mean(original_fragment, 0.05)
        weighted_mean_fragment = downscale_weighted_mean(original_fragment, 0.05)
        median_fragment = downscale_median(original_fragment, 0.05)

        memfile = image_plot_maker([original_fragment, nearest_neighbor_fragment, bilinear_fragment])
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('Oryginalny fragment | Metoda najbliższego sąsiada | Metoda interpolacji dwuliniowej')

        original_edges = cv2.Canny(original_fragment, threshold1=100, threshold2=200)
        nearest_neighbor_edges = cv2.Canny(nearest_neighbor_fragment, threshold1=100, threshold2=200)
        bilinear_edges = cv2.Canny(bilinear_fragment, threshold1=100, threshold2=200)

        memfile_edges = image_plot_maker([original_edges, nearest_neighbor_edges, bilinear_edges])
        document.add_picture(memfile_edges, width=Inches(6))
        memfile_edges.close()
        document.add_paragraph('Krawędzie: Oryginalny | Najbliższy sąsiad | Interpolacja dwuliniowa')

        memfile = image_plot_maker([mean_fragment, weighted_mean_fragment, median_fragment])
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('Średnia | Średnia ważona | Mediana')

        mean_edges = cv2.Canny(mean_fragment, threshold1=100, threshold2=200)
        weighted_mean_edges = cv2.Canny(weighted_mean_fragment, threshold1=100, threshold2=200)
        median_edges = cv2.Canny(median_fragment, threshold1=100, threshold2=200)

        memfile_edges = image_plot_maker([mean_edges, weighted_mean_edges, median_edges])
        document.add_picture(memfile_edges, width=Inches(6))
        memfile_edges.close()
        document.add_paragraph('Krawędzie: Średnia | Średnia ważona | Mediana')
        document.add_paragraph('')

document.save('report.docx')