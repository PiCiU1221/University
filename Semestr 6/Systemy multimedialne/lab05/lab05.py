import numpy as np
import matplotlib.pyplot as plt
from scipy.interpolate import interp1d
from docx import Document
from docx.shared import Inches
from io import BytesIO
import soundfile as sf
import scipy.fftpack

def Kwant(data, bit):
    d = (1 << bit) - 1
    DataF = data.astype(float)

    if np.issubdtype(data.dtype, np.floating):
        m = -1
        n = 1
    else:
        m = np.iinfo(data.dtype).min
        n = np.iinfo(data.dtype).max

    data_minus_m = DataF - m
    data_normalized = data_minus_m / (n - m)

    data_scaled = data_normalized * d
    data_quantized = np.round(data_scaled)
    data_divided_by_d = data_quantized / d

    data_rescaled = data_divided_by_d * (n - m)
    DataF = data_rescaled + m

    return DataF.astype(data.dtype)


def decimate(signal, n, original_fs):
    decimated_signal = signal[::n]

    new_sampling_rate = original_fs // n

    return decimated_signal, new_sampling_rate


def interpolate_signal(y, Fs, N1, use_linear=True):
    N = len(y)

    t = np.linspace(0, N / Fs, N, endpoint=False)
    t1 = np.linspace(0, N / Fs, N1, endpoint=False)

    if use_linear:
        metode = interp1d(t, y)
    else:
        metode = interp1d(t, y, kind='cubic')

    resampled_signal = metode(t1).astype(y.dtype)
    return resampled_signal

def plot_and_save_to_docx(doc, title, Signal, Fs, TimeMargin=[0, 0.02]):
    plt.figure(figsize=(12, 8))

    plt.subplot(2, 1, 1)
    plt.plot(np.arange(0, Signal.shape[0]) / Fs, Signal)
    plt.xlim(TimeMargin)
    plt.title("Sygnał")
    plt.xlabel("Czas (s)")
    plt.ylabel("Amplituda")

    fsize = 2 ** 8
    plt.subplot(2, 1, 2)
    yf = scipy.fftpack.fft(Signal, fsize)
    plt.plot(np.arange(0, Fs / 2, Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2]) + np.finfo(np.float32).eps))
    plt.title("Widmo")
    plt.xlabel("Częstotliwość (Hz)")
    plt.ylabel("Moc (dB)")

    plt.tight_layout()

    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)
    plt.close()

    doc.add_paragraph(title)
    doc.add_picture(img_stream, width=Inches(5))

def process_and_generate_plots():
    doc = Document()
    doc.add_heading('LAB06', 0)

    for file_name in ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav', 'sin_combined.wav']:
        doc.add_heading(file_name, 1)
        signal, fs = sf.read(f'SIN/{file_name}')

        for bit in [4, 8, 16, 24]:
            quantized_signal = Kwant(signal, bit)
            plot_and_save_to_docx(doc, f'Kwantyzacja {bit}-bit', quantized_signal, fs)

        for decimation_step in [2, 4, 6, 10, 24]:
            decimated_signal, new_fs = decimate(signal, decimation_step, fs)
            plot_and_save_to_docx(doc, f'Decymacja {decimation_step}', decimated_signal, new_fs)

        for fs_new in [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]:
            N1 = int(len(signal) * fs_new / fs)
            resampled_signal = interpolate_signal(signal, fs, N1, use_linear=True)
            plot_and_save_to_docx(doc, f'Interpolacja liniowa {fs_new} Hz', resampled_signal, fs_new)
            resampled_signal = interpolate_signal(signal, fs, N1, use_linear=False)
            plot_and_save_to_docx(doc, f'Interpolacja nieliniowa {fs_new} Hz', resampled_signal, fs_new)

    doc.save('report.docx')

def save_audio(signal, fs):
    signal = np.clip(signal, -1.0, 1.0)

    scaled_signal = np.int16(signal * 32767)

    filename = "temp_sound.wav"
    sf.write(filename, scaled_signal, fs)

def create_modified_audio_file(file_name, modification="quantization", bit_depth=8,
                                decimation_step=4, use_linear=True, new_fs=None):
    signal, fs = sf.read(file_name)

    if modification == "quantization":
        modified_signal = Kwant(signal, bit_depth)

    elif modification == "decimation":
        modified_signal, fs = decimate(signal, decimation_step, fs)

    elif modification == "interpolation":
        N1 = int(len(signal) * new_fs / fs)
        modified_signal = interpolate_signal(signal, fs, N1, use_linear)
        fs = new_fs

    save_audio(modified_signal, fs)

# generowanie plikow pierwszej czesci
# process_and_generate_plots()

# odwarzanie dzwiekow dla drugiej czesci
# create_modified_audio_file("SING/sing_low1.wav", modification="quantization", bit_depth=4)
# create_modified_audio_file("SING/sing_low1.wav", modification="quantization", bit_depth=8)
#
# create_modified_audio_file("SING/sing_low1.wav", modification="decimation", decimation_step=4)
# create_modified_audio_file("SING/sing_low1.wav", modification="decimation", decimation_step=6)
# create_modified_audio_file("SING/sing_low1.wav", modification="decimation", decimation_step=10)
# create_modified_audio_file("SING/sing_low1.wav", modification="decimation", decimation_step=24)
#
# create_modified_audio_file("SING/sing_low1.wav", modification="interpolation", new_fs=4000)
# create_modified_audio_file("SING/sing_low1.wav", modification="interpolation", new_fs=8000)
# create_modified_audio_file("SING/sing_low1.wav", modification="interpolation", new_fs=11999)
# create_modified_audio_file("SING/sing_low1.wav", modification="interpolation", new_fs=16000)
# create_modified_audio_file("SING/sing_low1.wav", modification="interpolation", new_fs=16953)
# create_modified_audio_file("SING/sing_low1.wav", modification="interpolation", new_fs=16953, use_linear=True)
