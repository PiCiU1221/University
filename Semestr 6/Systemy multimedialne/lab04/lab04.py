import numpy as np
import cv2
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO

pallet_1bit = np.linspace(0, 1, 2).reshape(2, 1)
pallet_2bit = np.linspace(0, 1, 4).reshape(4, 1)
pallet_4bit = np.linspace(0, 1, 16).reshape(16, 1)

pallet8 = np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 0.0, 1.0,],
        [0.0, 1.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [1.0, 0.0, 0.0,],
        [1.0, 0.0, 1.0,],
        [1.0, 1.0, 0.0,],
        [1.0, 1.0, 1.0,],
])

pallet16 =  np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [0.0, 0.0, 1.0,],
        [1.0, 0.0, 1.0,],
        [0.0, 0.5, 0.0,],
        [0.5, 0.5, 0.5,],
        [0.0, 1.0, 0.0,],
        [0.5, 0.0, 0.0,],
        [0.0, 0.0, 0.5,],
        [0.5, 0.5, 0.0,],
        [0.5, 0.0, 0.5,],
        [1.0, 0.0, 0.0,],
        [0.75, 0.75, 0.75,],
        [0.0, 0.5, 0.5,],
        [1.0, 1.0, 1.0,],
        [1.0, 1.0, 0.0,]
])

def colorFit(pixel, Pallet):
    distances = np.linalg.norm(Pallet - pixel, axis=1)
    closest_index = np.argmin(distances)
    return Pallet[closest_index]

def kwant_colorFit(img, Pallet):
    out_img = img.copy()
    for w in range(img.shape[0]):
        for k in range(img.shape[1]):
            out_img[w, k] = colorFit(img[w, k], Pallet)
    return out_img

# gray scale
# img = cv2.imread('IMG_GS/GS_0001.tif')
# img = cv2.imread('IMG_GS/GS_0003.png')
# img = cv2.imread('IMG_GS/GS_0002.png')
#
# img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
# img = img.astype(np.float32) / 255.0
#
# img_1bit = kwant_colorFit(img, pallet_1bit)
# img_2bit = kwant_colorFit(img, pallet_2bit)
# img_4bit = kwant_colorFit(img, pallet_4bit)
#
# fig, axs = plt.subplots(1, 4, figsize=(12, 4))
# axs[0].imshow(img, cmap="gray")
# axs[0].set_title("Original")
# axs[1].imshow(img_1bit, cmap="gray")
# axs[1].set_title("1-bit")
# axs[2].imshow(img_2bit, cmap="gray")
# axs[2].set_title("2-bit")
# axs[3].imshow(img_4bit, cmap="gray")
# axs[3].set_title("4-bit")
#
# plt.show()

# color
# img = cv2.imread('IMG_SMALL/SMALL_0009.jpg')
# img = cv2.imread('IMG_SMALL/SMALL_0007.jpg')
# img = cv2.imread('IMG_SMALL/SMALL_0004.jpg')
# img = cv2.imread('IMG_SMALL/SMALL_0006.jpg')

# img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
# img = img.astype(np.float32) / 255.0
#
# img_pallet8 = kwant_colorFit(img, pallet8)
# img_pallet16 = kwant_colorFit(img, pallet16)
#
# fig, axs = plt.subplots(1, 3, figsize=(12, 4))
# axs[0].imshow(img, cmap="gray")
# axs[0].set_title("Original")
# axs[1].imshow(img_pallet8, cmap="gray")
# axs[1].set_title("pallet8")
# axs[2].imshow(img_pallet16, cmap="gray")
# axs[2].set_title("pallet16")
#
# plt.show()

def random_dithering(img):
    random_matrix = np.random.rand(img.shape[0], img.shape[1])
    result_img = (img >= random_matrix).astype(np.uint8)
    return result_img

# img = cv2.imread('IMG_GS/GS_0001.tif', cv2.IMREAD_GRAYSCALE)
# img = img.astype(np.float32) / 255.0
#
# binary_img = random_dithering(img)
# plt.imshow(binary_img, cmap='gray')
# plt.show()

M1 = np.array([
    [0, 2],
    [3, 1]])

M2 = np.array([
    [0,  8,  2, 10],
    [12, 4, 14, 6],
    [3, 11, 1,  9],
    [15, 7, 13, 5]
])

M4 = np.array([
    [ 0, 32,  8, 40,  2, 34, 10, 42],
    [48, 16, 56, 24, 50, 18, 58, 26],
    [12, 44,  4, 36, 14, 46,  6, 38],
    [60, 28, 52, 20, 62, 30, 54, 22],
    [ 3, 35, 11, 43,  1, 33,  9, 41],
    [51, 19, 59, 27, 49, 17, 57, 25],
    [15, 47,  7, 39, 13, 45,  5, 37],
    [63, 31, 55, 23, 61, 29, 53, 21]
])

def ordered_dithering(img, palette, M):
    n = M.shape[0]
    r = 1

    Mpre = (M + 1) / n ** 2 - 0.5

    out_img = img.copy()

    for y in range(img.shape[0]):
        for x in range(img.shape[1]):
            new_color = img[y, x] + r * Mpre[y % n, x % n]
            out_img[y, x] = colorFit(new_color, palette)

    return out_img

# img = cv2.imread('IMG_SMALL/SMALL_0002.png')
# img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
# img = img.astype(np.float32) / 255.0
#
# result_img = ordered_dithering(img, pallet_1bit, M2)
#
# plt.imshow(result_img, cmap='gray')
# plt.show()

def floyd_steinberg_dithering(img, palette):
    img_dithered = img.copy()
    h, w = img_dithered.shape[:2]

    for y in range(h):
        for x in range(w):
            old_pixel = img_dithered[y, x].copy()
            new_pixel = colorFit(old_pixel, palette)
            img_dithered[y, x] = new_pixel
            quant_error = old_pixel - new_pixel

            if x + 1 < w:
                img_dithered[y, x + 1] += quant_error * 7 / 16
            if y + 1 < h and x > 0:
                img_dithered[y + 1, x - 1] += quant_error * 3 / 16
            if y + 1 < h:
                img_dithered[y + 1, x] += quant_error * 5 / 16
            if y + 1 < h and x + 1 < w:
                img_dithered[y + 1, x + 1] += quant_error * 1 / 16

    return np.clip(img_dithered, 0.0, 1.0)

# img = cv2.imread('IMG_SMALL/SMALL_0002.png')
# img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
# img = img.astype(np.float32) / 255.0
#
# result_img = floyd_steinberg_dithering(img, pallet_1bit)
#
# plt.imshow(result_img, cmap='gray')
# plt.show()

# first plot type

# img = cv2.imread('IMG_GS/GS_0002.png')
#
# img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
# img = img.astype(np.float32) / 255.0
#
# quantized_img = kwant_colorFit(img, pallet_1bit)
# ordered_dithered_img = ordered_dithering(img, pallet_1bit, M2)
# random_dithered_img = random_dithering(img)
# floyd_steinberg_dithered_img = floyd_steinberg_dithering(img, pallet_1bit)
#
# fig, axes = plt.subplots(2, 3, figsize=(12, 8))
#
# axes[0, 0].imshow(img, cmap='gray')
# axes[0, 0].set_title("Oryginał")
# axes[0, 0].axis('off')
#
# axes[0, 1].imshow(quantized_img, cmap='gray')
# axes[0, 1].set_title("Kwantyzacja")
# axes[0, 1].axis('off')
#
# axes[0, 2].imshow(ordered_dithered_img, cmap='gray')
# axes[0, 2].set_title("Dithering Zorganizowany")
# axes[0, 2].axis('off')
#
# axes[1, 1].imshow(random_dithered_img, cmap='gray')
# axes[1, 1].set_title("Dithering Losowy")
# axes[1, 1].axis('off')
#
# axes[1, 2].imshow(floyd_steinberg_dithered_img, cmap='gray')
# axes[1, 2].set_title("Dithering Floyda-Steinberga")
# axes[1, 2].axis('off')
#
# axes[1, 0].axis('off')
#
# fig.suptitle("Dithering 1-bit", fontsize=16)
#
# plt.tight_layout()
# plt.subplots_adjust(top=0.9)
#
# plt.show()

# second plot type

# img = cv2.imread('IMG_GS/GS_0002.png')
# img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
# img = img.astype(np.float32) / 255.0
#
# quantized_img = kwant_colorFit(img, pallet_2bit)
# ordered_dithered_img = ordered_dithering(img, pallet_2bit, M2)
# floyd_steinberg_dithered_img = floyd_steinberg_dithering(img, pallet_2bit)
#
# fig, axes = plt.subplots(2, 2, figsize=(10, 8))
#
# axes[0, 0].imshow(img, cmap='gray')
# axes[0, 0].set_title("Oryginał")
# axes[0, 0].axis('off')
#
# axes[0, 1].imshow(ordered_dithered_img, cmap='gray')
# axes[0, 1].set_title("Dithering Zorganizowany")
# axes[0, 1].axis('off')
#
# axes[1, 0].imshow(quantized_img, cmap='gray')
# axes[1, 0].set_title("Kwantyzacja")
# axes[1, 0].axis('off')
#
# axes[1, 1].imshow(floyd_steinberg_dithered_img, cmap='gray')
# axes[1, 1].set_title("Dithering Floyda-Steinberga")
# axes[1, 1].axis('off')
#
# fig.suptitle("Dithering 2-bity", fontsize=16)
#
# plt.tight_layout()
# plt.show()

def save_dithered_plots(img, quantized_img, ordered_dithered_img,
                        floyd_steinberg_dithered_img, title, doc):
    fig, axes = plt.subplots(2, 2, figsize=(10, 8))

    axes[0, 0].imshow(img, cmap='gray')
    axes[0, 0].set_title("Oryginał")
    axes[0, 0].axis('off')

    axes[0, 1].imshow(ordered_dithered_img, cmap='gray')
    axes[0, 1].set_title("Dithering Zorganizowany")
    axes[0, 1].axis('off')

    axes[1, 0].imshow(quantized_img, cmap='gray')
    axes[1, 0].set_title("Kwantyzacja")
    axes[1, 0].axis('off')

    axes[1, 1].imshow(floyd_steinberg_dithered_img, cmap='gray')
    axes[1, 1].set_title("Dithering Floyda-Steinberga")
    axes[1, 1].axis('off')

    fig.suptitle(title, fontsize=16)

    img_stream = BytesIO()
    plt.tight_layout()
    plt.savefig(img_stream, format='png')
    plt.close(fig)
    img_stream.seek(0)

    doc.add_picture(img_stream, width=Inches(6))
    doc.add_page_break()

gray_image_names = ['IMG_GS/GS_0002.png', 'IMG_GS/GS_0001.tif', 'IMG_GS/GS_0003.png']
color_image_names = ['IMG_SMALL/SMALL_0004.jpg', 'IMG_SMALL/SMALL_0006.jpg', 'IMG_SMALL/SMALL_0007.jpg', 'IMG_SMALL/SMALL_0009.jpg']

doc = Document()

for img_path in gray_image_names:
    doc.add_paragraph(f"Obraz: {img_path}")

    img = cv2.imread(img_path)
    img_gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    img_gray = img_gray.astype(np.float32) / 255.0

    quantized_img_1bit = kwant_colorFit(img_gray, pallet_1bit)
    ordered_dithered_img_1bit = ordered_dithering(img_gray, pallet_1bit, M2)
    random_dithered_img_1bit = random_dithering(img_gray)
    floyd_steinberg_dithered_img_1bit = floyd_steinberg_dithering(img_gray, pallet_1bit)

    quantized_img_2bit = kwant_colorFit(img_gray, pallet_2bit)
    ordered_dithered_img_2bit = ordered_dithering(img_gray, pallet_2bit, M2)
    floyd_steinberg_dithered_img_2bit = floyd_steinberg_dithering(img_gray, pallet_2bit)

    quantized_img_4bit = kwant_colorFit(img_gray, pallet_4bit)
    ordered_dithered_img_4bit = ordered_dithering(img_gray, pallet_4bit, M2)
    floyd_steinberg_dithered_img_4bit = floyd_steinberg_dithering(img_gray, pallet_4bit)

    fig, axes = plt.subplots(2, 3, figsize=(12, 8))

    axes[0, 0].imshow(img_gray, cmap='gray')
    axes[0, 0].set_title("Oryginał")
    axes[0, 0].axis('off')

    axes[0, 1].imshow(quantized_img_1bit, cmap='gray')
    axes[0, 1].set_title("Kwantyzacja")
    axes[0, 1].axis('off')

    axes[0, 2].imshow(ordered_dithered_img_1bit, cmap='gray')
    axes[0, 2].set_title("Dithering Zorganizowany")
    axes[0, 2].axis('off')

    axes[1, 1].imshow(random_dithered_img_1bit, cmap='gray')
    axes[1, 1].set_title("Dithering Losowy")
    axes[1, 1].axis('off')

    axes[1, 2].imshow(floyd_steinberg_dithered_img_1bit, cmap='gray')
    axes[1, 2].set_title("Dithering Floyda-Steinberga")
    axes[1, 2].axis('off')

    axes[1, 0].axis('off')

    fig.suptitle("Dithering 1-bit", fontsize=16)

    img_stream_1bit = BytesIO()
    plt.subplots_adjust(left=0, right=1, top=0.90, bottom=0.05, wspace=-0.4, hspace=0.1)
    plt.savefig(img_stream_1bit, format='png', bbox_inches='tight', pad_inches=0.1)
    plt.close(fig)
    img_stream_1bit.seek(0)

    doc.add_picture(img_stream_1bit, width=Inches(6))
    doc.add_page_break()

    save_dithered_plots(img, quantized_img_2bit, ordered_dithered_img_2bit,
                        floyd_steinberg_dithered_img_2bit, "Dithering 2-bity", doc)

    save_dithered_plots(img, quantized_img_4bit, ordered_dithered_img_4bit,
                        floyd_steinberg_dithered_img_4bit, "Dithering 4-bity", doc)

for img_path in color_image_names:
    doc.add_paragraph(f"Obraz: {img_path}")

    img = cv2.imread(img_path)
    img_color = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    img_color = img_color.astype(np.float32) / 255.0

    quantized_img_pallet8 = kwant_colorFit(img_color, pallet8)
    ordered_dithered_img_pallet8 = ordered_dithering(img_color, pallet8, M2)
    floyd_steinberg_dithered_img_pallet8 = floyd_steinberg_dithering(img_color, pallet8)

    quantized_img_pallet16 = kwant_colorFit(img_color, pallet16)
    ordered_dithered_img_pallet16 = ordered_dithering(img_color, pallet16, M2)
    floyd_steinberg_dithered_img_pallet16 = floyd_steinberg_dithering(img_color, pallet16)

    save_dithered_plots(img_color, quantized_img_pallet8, ordered_dithered_img_pallet8,
                        floyd_steinberg_dithered_img_pallet8, "Dithering paleta 8 kolorów", doc)

    save_dithered_plots(img_color, quantized_img_pallet16, ordered_dithered_img_pallet16,
                        floyd_steinberg_dithered_img_pallet16, "Dithering paleta 16 kolorów", doc)

doc.save("report.docx")
