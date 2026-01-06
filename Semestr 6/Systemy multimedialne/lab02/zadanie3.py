import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

from Lab1_img import image_plot_maker

df = pd.DataFrame()

df = pd.DataFrame(
    data={
        'Filename':['IMG_INTRO/B01.png', 'IMG_INTRO/B02.jpg'],'Grayscale':[False, False],
        'Fragments':[[[50,200,250,400],[300,1200,500,1400]], [[350,300,550,500],[550,500,750,700]]]
    })

document = Document()

for index, row in df.iterrows():
    img = plt.imread(row['Filename'])
    # if row['Grayscale']:
    #     # GS image - teraz nas nie intersuje
    # else:
    #     # Obraz kolowowy
    if row['Fragments'] is not None:
        # mamy nie pustą listę fragmentów
        for f in row['Fragments']:
            fragment = img[f[0]:f[2],f[1]:f[3]].copy()
            # tu wykonujesz operacje i inne wyświetlenia na fragmencie

            document.add_heading('Plik - {}'.format(row["Filename"]), 2)

            memfile = image_plot_maker(fragment)

            document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

            memfile.close()
            document.add_paragraph(f'Fragment: {f}')

document.save('report.docx')