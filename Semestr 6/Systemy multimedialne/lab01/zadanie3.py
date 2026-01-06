from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
import soundfile as sf
import scipy.fftpack

def plotAudio(Signal, Fs, fsize, axs):
    time = np.arange(0, Signal.shape[0]) / Fs
    axs[0].plot(time, Signal)
    axs[0].set_xlim([0, 0.02])
    axs[0].set_title("Signal in Time Domain")
    axs[0].set_xlabel("Time (s)")
    axs[0].set_ylabel("Amplitude")

    yf = scipy.fftpack.fft(Signal, fsize)
    freqs = np.arange(0, Fs / 2, Fs / fsize)
    magnitude = np.abs(yf[:fsize // 2])
    magnitude = 20 * np.log10(magnitude)
    axs[1].plot(freqs, magnitude)
    axs[1].set_title("Frequency Spectrum (fsize={})".format(fsize))
    axs[1].set_xlabel("Frequency (Hz)")
    axs[1].set_ylabel("Magnitude (dB)")

    peak_freq = freqs[np.argmax(magnitude)]
    peak_magnitude = magnitude[np.argmax(magnitude)]
    return peak_freq, peak_magnitude

document = Document()

files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
fsize_list=[2**8,2**12,2**16]
for file in files:
    document.add_heading('Plik - {}'.format(file), 2)
    data, fs = sf.read(f"SIN/{file}", dtype=np.int32)

    for i, fsize in enumerate(fsize_list):
        document.add_heading('fsize {}'.format(fsize), 3)  # nagłówek sekcji, mozę być poziom wyżej
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))  # tworzenie plota

        ############################################################
        # Tu wykonujesz jakieś funkcje i rysujesz wykresy
        ############################################################

        peak_freq, peak_magnitude = plotAudio(data, fs, fsize, axs)

        # fig.suptitle('Time margin {}'.format(Margin))  # Tytuł wykresu
        fig.tight_layout(pad=1.5)  # poprawa czytelności
        memfile = BytesIO()  # tworzenie bufora
        fig.savefig(memfile)  # z zapis do bufora

        document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

        memfile.close()
        ############################################################
        # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
        document.add_paragraph(f'Najwyższa wartość widma: {peak_magnitude:.2f} dB przy częstotliwości {peak_freq:.2f} Hz')
        ############################################################

document.save('report.docx')  # zapis do pliku