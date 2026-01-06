import os
import cv2
import numpy as np
from skimage.metrics import structural_similarity as ssim
from docx import Document
from docx.shared import Inches

def mse(modified, original):
    modified = modified.astype(np.float64)
    original = original.astype(np.float64)
    m, n = original.shape

    total_error = 0.0
    for i in range(m):
        for j in range(n):
            diff = (modified[i, j] - original[i, j]) ** 2
            total_error += diff

    mse_val = total_error / (m * n)
    return mse_val

def nmse(modified, original):
    zero_img = np.zeros_like(modified, dtype=np.float64)
    numerator = mse(modified, original)
    denominator = mse(modified, zero_img)

    return numerator / denominator

def psnr(modified, original):
    max_pixel = np.iinfo(np.uint8).max

    mse_val = mse(modified, original)
    psnr_val = 10 * np.log10((max_pixel ** 2) / mse_val)

    return psnr_val

def image_fidelity(modified, original):
    modified = modified.astype(np.float64)
    original = original.astype(np.float64)
    m, n = original.shape

    numerator = 0.0
    denominator = 0.0
    for i in range(m):
        for j in range(n):
            numerator += (modified[i, j] - original[i, j]) ** 2
            denominator += (modified[i, j] * original[i, j])

    fidelity = 1 - (numerator / denominator)
    return fidelity

def compute_metrics(modified, original):
    modified_gray = cv2.cvtColor(modified, cv2.COLOR_RGB2GRAY)
    original_gray = cv2.cvtColor(original, cv2.COLOR_RGB2GRAY)

    mse_val = mse(modified_gray, original_gray)
    nmse_val = nmse(modified_gray, original_gray)
    psnr_val = psnr(modified_gray, original_gray)
    fidelity_val = image_fidelity(modified_gray, original_gray)
    ssim_val = ssim(original_gray, modified_gray, win_size=7, data_range=255)

    return mse_val, nmse_val, psnr_val, fidelity_val, ssim_val

input_dir = './images'
output_dir = './output'
os.makedirs(output_dir, exist_ok=True)

images = ['zdjecie1.png', 'zdjecie2.png', 'zdjecie3.png', 'zdjecie4.png']
methods = ['Kompresja JPEG', 'Rozmycie Gaussowskie', 'Rozmycie medianowe', 'Szum Gaussowski']

doc = Document()
doc.add_heading('Ocena jakości obrazów', 0)

for idx, filename in enumerate(images):
    image_path = os.path.join(input_dir, filename)
    image = cv2.imread(image_path)
    image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)

    doc.add_heading(f'{filename} - {methods[idx]}', level=1)

    orig_path = os.path.join(output_dir, f'orig_{filename}')
    cv2.imwrite(orig_path, image)
    doc.add_paragraph("Obraz oryginalny")
    doc.add_picture(orig_path, width=Inches(3.5))

    metrics_table = [['Parametr', 'MSE', 'NMSE', 'PSNR', 'IF', 'SSIM']]

    if idx == 0:
        for q in [95, 80, 60, 40, 20]:
            degraded_path = os.path.join(output_dir, f'{filename}_jpeg_q{q}.jpg')
            cv2.imwrite(degraded_path, image, [cv2.IMWRITE_JPEG_QUALITY, q])
            degraded = cv2.imread(degraded_path)
            degraded_rgb = cv2.cvtColor(degraded, cv2.COLOR_BGR2RGB)

            mse_val, nmse_val, psnr_val, fidelity_val, ssim_val = compute_metrics(image_rgb, degraded_rgb)

            doc.add_paragraph(f'Jakość kompresji JPEG = {q}')
            doc.add_picture(degraded_path, width=Inches(3.5))
            metrics_table.append(
                [f'Q={q}', f'{mse_val:.2f}', f'{nmse_val:.4f}', f'{psnr_val:.2f}', f'{fidelity_val:.4f}', f'{ssim_val:.4f}'])

    elif idx == 1:
        for k in [3, 5, 7, 9, 11]:
            degraded = cv2.GaussianBlur(image, (k, k), sigmaX=0, sigmaY=0)
            degraded_path = os.path.join(output_dir, f'{filename}_gauss_k{k}.png')
            cv2.imwrite(degraded_path, degraded)

            degraded_rgb = cv2.cvtColor(degraded, cv2.COLOR_BGR2RGB)
            mse_val, nmse_val, psnr_val, fidelity_val, ssim_val = compute_metrics(image_rgb, degraded_rgb)

            doc.add_paragraph(f'Rozmycie Gaussowskie, rozmiar maski = {k}, sigmaX = 0, sigmaY = 0')
            doc.add_picture(degraded_path, width=Inches(3.5))
            metrics_table.append(
                [f'k={k}', f'{mse_val:.2f}', f'{nmse_val:.4f}', f'{psnr_val:.2f}', f'{fidelity_val:.4f}', f'{ssim_val:.4f}'])

    elif idx == 2:
        for k in [3, 5, 7, 9, 11]:
            degraded = cv2.medianBlur(image, k)
            degraded_path = os.path.join(output_dir, f'{filename}_median_k{k}.png')
            cv2.imwrite(degraded_path, degraded)

            degraded_rgb = cv2.cvtColor(degraded, cv2.COLOR_BGR2RGB)
            mse_val, nmse_val, psnr_val, fidelity_val, ssim_val = compute_metrics(image_rgb, degraded_rgb)

            doc.add_paragraph(f'Rozmycie medianowe, rozmiar maski = {k}')
            doc.add_picture(degraded_path, width=Inches(3.5))
            metrics_table.append(
                [f'k={k}', f'{mse_val:.2f}', f'{nmse_val:.4f}', f'{psnr_val:.2f}', f'{fidelity_val:.4f}', f'{ssim_val:.4f}'])

    elif idx == 3:
        sigma = 25
        alpha_values = [0.1, 0.2, 0.3, 0.5, 0.7]

        for alpha in alpha_values:
            row, col, ch = image.shape

            gauss = np.random.normal(0, sigma, (row, col, ch))

            noisy = (image.astype(np.float64) + alpha * gauss).clip(0, 255).astype(np.uint8)

            degraded_path = os.path.join(output_dir, f'{filename}_gaussNoise_a{alpha:.3f}.png')
            cv2.imwrite(degraded_path, noisy)

            noisy_rgb = cv2.cvtColor(noisy, cv2.COLOR_BGR2RGB)
            mse_val, nmse_val, psnr_val, fidelity_val, ssim_val = compute_metrics(image_rgb, noisy_rgb)

            doc.add_paragraph(f'Szum Gaussowski alpha = {alpha:.3f}')
            doc.add_picture(degraded_path, width=Inches(3.5))
            metrics_table.append(
                [f'a={alpha:.3f}', f'{mse_val:.2f}', f'{nmse_val:.4f}', f'{psnr_val:.2f}', f'{fidelity_val:.4f}', f'{ssim_val:.4f}'])

    doc.add_paragraph("\nTabela miar jakości obrazu:")
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(metrics_table[0]):
        hdr_cells[i].text = header

    for row in metrics_table[1:]:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val

doc.save('raport.docx')
